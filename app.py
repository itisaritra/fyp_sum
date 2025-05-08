import streamlit as st
import cohere
import time
import PyPDF2
import docx
from io import StringIO, BytesIO

# Initialize Cohere client
API_KEY = "ms0slgGuFh2udoxiN8zIYPK7RsHRbqU03IQGpBpB"  # Replace with your actual API key
co = cohere.Client(api_key=API_KEY)

# Define Cohere summarization
def generate_summary_cohere(text, length_option):
    length_mapping = {
        "Short": "short",
        "Medium": "medium",
        "Long": "long"
    }
    length = length_mapping[length_option]

    try:
        # Make the API call to Cohere's summarize endpoint
        response = co.summarize(
            text=text,
            length=length,
            format="paragraph",
            model="summarize-xlarge",
            additional_command="",
            temperature=0.3
        )
        time.sleep(1)  # Simulating delay
        return response.summary
    except Exception as e:
        return f"An error occurred with Cohere: {e}"

# Extract text from uploaded files
def extract_text(file):
    if file.type == "text/plain":  # .txt file
        return StringIO(file.getvalue().decode("utf-8")).read()
    elif file.type == "application/pdf":  # .pdf file
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":  # .docx file
        doc = docx.Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text
        return text
    else:
        return None

# Streamlit App
def main():
    st.set_page_config(page_title="Ensemble Summarization App", page_icon="📝", layout="wide")

    # Header Section with Styling
    st.markdown(
        """
        <style>
            .main-header {
                font-size: 2.5em;
                font-weight: bold;
                color: #4CAF50;
                text-align: center;
                margin-bottom: 20px;
            }
            .sub-header {
                text-align: center;
                font-size: 1.2em;
                color: #6c757d;
                margin-bottom: 30px;
            }
        </style>
        <div class="main-header">Ensemble Text Summarization</div>
        <div class="sub-header">
            A Streamlit-based application designed to enhance text summarization accuracy using an advanced ensemble learning approach. The architecture integrates fine-tuned BART and Pegasus models, trained on a curated dataset of 3 lakh article-summary pairs, with RoBERTa as the meta-model for improved performance. This combination leverages the strengths of multiple models to generate precise, coherent, and contextually accurate summaries, accessible through an intuitive and interactive user interface.
        </div>
        """, unsafe_allow_html=True
    )

    # Sidebar Section
    st.sidebar.header("⚙️ Options")
    length_option = st.sidebar.selectbox("Choose Summary Length:", ("Short", "Medium", "Long"))

    st.sidebar.markdown("---")
    st.sidebar.markdown("### 📄 Resources")
    st.sidebar.markdown("- [Project Report](https://drive.google.com/file/d/12iWuXNXkGBNk5PrK6mHJO6RL5Ohwa5jS/view)")
    st.sidebar.markdown("- [Project Presentation](https://summari-wa22e81.gamma.site/)")

    st.sidebar.markdown("---")
    st.sidebar.write(
        """
        Final Year Project Developed by:  
        - [Aritra Ghosh](https://aritraghosh.co/) (B.Tech CSE) (Section A-74) (12021002002137)  
        - [Subhojit Ghosh](https://subhojit.pages.dev/) (B.Tech CSE) (Section B-97) (12021002002160)  
        
        Under the mentorship of Dr. Anupam Mondal.
        """
    )

    # File Upload Section
    st.markdown("### 📂 Upload a File or Enter Text for Summarization")
    uploaded_file = st.file_uploader("Upload a .txt, .docx, or .pdf file:", type=["txt", "docx", "pdf"])
    text = None
    if uploaded_file:
        text = extract_text(uploaded_file)
        if text:
            st.success("✅ File uploaded successfully!")
        else:
            st.error("🚫 Unsupported file type or error reading the file.")

    # Text Area Section
    if not text:
        text = st.text_area("Or enter your text here:", placeholder="Type or paste text here...", height=200)

    if st.button("📝 Generate Summary"):
        if text:
            with st.spinner("Processing... Please wait."):
                # Show a progress bar
                progress = st.progress(0)
                for percent in range(1, 101):
                    time.sleep(0.05)  # Simulating progress
                    progress.progress(percent)

                # Generate summary
                summary = generate_summary_cohere(text, length_option)

                # Display summary
                st.markdown("### 📝 Generated Summary")
                st.text_area("Summary:", summary, height=200)

                # Download button
                st.markdown("---")
                st.markdown("### 💾 Download Summary")
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="Download as .txt",
                        data=summary,
                        file_name="summary.txt",
                        mime="text/plain"
                    )
                with col2:
                    docx_file = BytesIO()  # Use BytesIO for binary stream
                    doc = docx.Document()
                    doc.add_paragraph(summary)
                    doc.save(docx_file)
                    docx_file.seek(0)  # Reset the stream position to the start
                    st.download_button(
                        label="Download as .docx",
                        data=docx_file.getvalue(),  # Use .getvalue() to get the binary content
                        file_name="summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.error("🚫 Please provide some text or upload a file.")

    # Footer Section
    st.markdown("---")
    st.markdown(
        """
        <style>
            .footer {
                text-align: center;
                font-size: 0.9em;
                color: #6c757d;
            }
        </style>
        <div class="footer">
            © 2024 Multi-Model Summarization | All rights reserved.
        </div>
        """, unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()
